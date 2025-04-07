from typing import Dict, Any, List, Optional
import logging
import asyncio
import uuid
import os
from datetime import datetime
from docx import Document
from sqlalchemy.orm import Session
from app.core.database import get_db
from app.db.models import Message, Document as DocumentModel
from app.core.security import generate_share_token
from app.agents.document_prep import document_prep_agent
from app.core.config import settings

logger = logging.getLogger(__name__)

class DocumentService:
    """
    Service for generating and managing legal documents
    """
    
    async def generate_document(self, message_id: uuid.UUID, query: str, 
                          parameters: Dict[str, Any], db: Session) -> Dict[str, Any]:
        """
        Generate a legal document based on a query and parameters
        """
        # Create document record in database
        document = DocumentModel(
            message_id=message_id,
            title=f"Документ от {datetime.now().strftime('%d.%m.%Y %H:%M')}",
            content="", # Initially empty
            document_type="pending",
            share_token=generate_share_token()
        )
        db.add(document)
        db.commit()
        db.refresh(document)
        
        # Start async document generation
        asyncio.create_task(self._generate_document_async(
            message_id,
            document.id,
            query,
            parameters,
            db
        ))
        
        return {
            "document_id": document.id,
            "share_token": document.share_token,
            "status": "generating"
        }
    
    async def _generate_document_async(self, message_id: uuid.UUID, document_id: uuid.UUID,
                                 query: str, parameters: Dict[str, Any], db: Session) -> None:
        """
        Asynchronously generate a document
        """
        db_for_task = next(get_db())  # Get a new DB session for this async task
        
        try:
            # Generate document
            result = await document_prep_agent.process_query(query, parameters)
            
            if "error" in result:
                # Update document with error
                document = db_for_task.query(DocumentModel).filter(DocumentModel.id == document_id).first()
                document.content = f"Произошла ошибка при генерации документа: {result['error']}"
                document.document_type = "error"
                db_for_task.commit()
                return
            
            document_content = result.get("document", "")
            document_type = result.get("document_type", "unknown")
            
            # Update document
            document = db_for_task.query(DocumentModel).filter(DocumentModel.id == document_id).first()
            document.content = document_content
            document.document_type = document_type
            document.title = f"{document_type.capitalize()} от {datetime.now().strftime('%d.%m.%Y %H:%M')}"
            db_for_task.commit()
            
            # Also update the related message
            message = db_for_task.query(Message).filter(Message.id == message_id).first()
            if message:
                message.content += f"\n\n**Документ сгенерирован!** [Скачать документ](/api/v1/documents/{document.share_token}/download)"
                db_for_task.commit()
            
        except Exception as e:
            logger.error(f"Error generating document: {str(e)}")
            
            # Update document with error
            document = db_for_task.query(DocumentModel).filter(DocumentModel.id == document_id).first()
            document.content = f"Произошла ошибка при генерации документа: {str(e)}"
            document.document_type = "error"
            db_for_task.commit()
        finally:
            db_for_task.close()
    
    def get_document(self, document_id: uuid.UUID, db: Session) -> Optional[Dict[str, Any]]:
        """
        Get document by ID
        """
        document = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
        
        if not document:
            return None
        
        return {
            "id": str(document.id),
            "title": document.title,
            "content": document.content,
            "document_type": document.document_type,
            "share_token": document.share_token,
            "created_at": document.created_at.isoformat()
        }
    
    def get_document_by_token(self, share_token: str, db: Session) -> Optional[Dict[str, Any]]:
        """
        Get document by share token
        """
        document = db.query(DocumentModel).filter(DocumentModel.share_token == share_token).first()
        
        if not document:
            return None
        
        return {
            "id": str(document.id),
            "title": document.title,
            "content": document.content,
            "document_type": document.document_type,
            "share_token": document.share_token,
            "created_at": document.created_at.isoformat()
        }
    
    def create_docx(self, document_id: uuid.UUID, db: Session) -> Optional[str]:
        """
        Create a DOCX file from a document
        """
        document = db.query(DocumentModel).filter(DocumentModel.id == document_id).first()
        
        if not document:
            return None
        
        # Create DOCX file
        docx = Document()
        
        # Add title
        docx.add_heading(document.title, 0)
        
        # Add content
        for paragraph in document.content.split('\n'):
            if paragraph.strip():
                docx.add_paragraph(paragraph)
        
        # Ensure uploads directory exists
        os.makedirs(settings.UPLOAD_DIR, exist_ok=True)
        
        # Save file
        file_path = os.path.join(settings.UPLOAD_DIR, f"{document.share_token}.docx")
        docx.save(file_path)
        
        return file_path

def is_valid_docx(file_path):
    """
    Проверяет, является ли файл корректным DOCX (ZIP-архивом).
    """
    try:
        with open(file_path, "rb") as f:
            header = f.read(4)
        return header == b'PK\x03\x04'  # DOCX - это zip-архив, который начинается с PK
    except Exception:
        return False

def extract_text_from_docx(docx_file_path):
    """
    Извлекает текст из .docx и удаляет пустые строки.
    """
    if not os.path.exists(docx_file_path):
        raise ValueError(f"Файл не найден: {docx_file_path}")

    # Проверяем, является ли файл ZIP-архивом (DOCX)
    if not is_valid_docx(docx_file_path):
        raise ValueError(f"Ошибка: файл {docx_file_path} не является корректным DOCX!")

    try:
        document = Document(docx_file_path)
        full_text = "\n".join([p.text.strip() for p in document.paragraphs if p.text.strip()])
        if full_text.strip():
            return full_text
    except Exception as e:
        logger.error(f"Ошибка при обработке python-docx: {e}")
        raise ValueError(f"Ошибка при обработке DOCX: {e}")

def extract_text_from_any_document(file_path: str) -> str:
    """
    Универсальная функция для извлечения текста из документов.
    Пока поддерживает только DOCX.
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Неподдерживаемый тип файла: {ext}. В данный момент поддерживаются только файлы DOCX.")

# Create singleton instance
document_service = DocumentService()